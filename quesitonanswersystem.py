# -*- coding: utf-8 -*-
"""QuesitonAnswerSystem.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1Re9hgtt3ytajPeXkAJ290mj51_4lIs0A

# **Requirements**
"""

# !pip install deeppavlov
# !python -m deeppavlov install squad_bert
!pip install rake-nltk
!pip install fuzzywuzzy
!pip install python-docx

from deeppavlov import build_model, configs
model_qa = build_model(configs.squad.squad_bert, download=True)

import docx 
import nltk
nltk.download('averaged_perceptron_tagger')
from rake_nltk import Rake, Metric
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from fuzzywuzzy import fuzz
from fuzzywuzzy import process
from nltk.stem import PorterStemmer
from nltk.stem.snowball import SnowballStemmer

"""# **Functions**"""

rake = Rake(min_length=1, max_length=3, ranking_metric=Metric.DEGREE_TO_FREQUENCY_RATIO)
doc = docx.Document('/content/drive/My Drive/SIH.docx')

def pos_tag_check(pos_tag):
    for word,tag in pos_tag:
        if(pos_tag.index((word,tag))<5):
            if(tag in ['PRP', 'PRP$', 'RB', 'RBR', 'RBS', 'WRB', 'WDT', 'WP', 'WP$', 'DT']):#follows/given as
                return 1
                break
        else:
            break
    return 0

def removeTags(pos_tag):
    a=''
    for word,tag in pos_tag:
        if(tag in ['NNS' ,'RB']):
            continue
        elif word in stop_words:
            continue
        else:
            if(a==''):
                a= word
            else:
                a = a + ' ' + word
    return a


def match(t1,t2,parti=80,rati=60,sortRati=65):
    for x in t1:
        for y in t2:
            count=0
            if(fuzz.partial_ratio(x,y) > parti):
                count+=1
            if(fuzz.ratio(x,y) > rati):
                count+=1
            if(fuzz.token_sort_ratio(x,y)>sortRati):
                count+=1
            if(count>=3):
                return 1
    return 0


def checkKeywordMatch(first_sent,second_sent):
    t1=[]
    t2=[]
    r = Rake(min_length=1, max_length=3)
    r.extract_keywords_from_text(first_sent)
    a= r.get_ranked_phrases()
    for b in a:
        b=b.split()
        pos_tag = nltk.pos_tag(b)
        temp=removeTags(pos_tag)
        if(temp != ''):
            t1.append(removeTags(pos_tag).lower())
    r.extract_keywords_from_text(second_sent)
    c= r.get_ranked_phrases()
    for d in c:
        d=d.split()
        pos_tag = nltk.pos_tag(d)
        temp =removeTags(pos_tag)
        if(temp != ''):
            t2.append(removeTags(pos_tag).lower())
    if(match(t1,t2)):
        return 1
    else:
        return 0

def clustersCreater(data):
    clustSent=''
    keyW = []
    finalData = []
    sent_text = nltk.sent_tokenize(data) 
    for s in sent_text:
        flag=0
        tokenized_text = nltk.word_tokenize(s)
        if(sent_text.index(s) != 0):
            if len(tokenized_text) < 4:
                finalData.append([keyW,clustSent])
                keyW=[]
                clustSent=""
                x = s.split()
                pos_tag = nltk.pos_tag(x)
                keyW.append(removeTags(pos_tag))

                testFlag =1
            else:
                pos_tag = nltk.pos_tag(s.split())
                pos_tagged_noun_verb = []
                if(pos_tag_check(pos_tag)):
                    flag = 1
                if(flag==1 or testFlag == 1):
                    testFlag = 0
                    rake.extract_keywords_from_text(s)
                    for x in rake.get_ranked_phrases():
                        x = x.split()
                        pos_tag = nltk.pos_tag(x)
                        keyW.append(removeTags(pos_tag))                     
                    clustSent = clustSent + s
                else:
                    first_sent = sent_text[sent_text.index(s)-1]
                    second_sent = s
                    if checkKeywordMatch(first_sent,second_sent):
                        rake.extract_keywords_from_text(s)
                        for x in rake.get_ranked_phrases():
                            x = x.split()
                            pos_tag = nltk.pos_tag(x)
                            keyW.append(removeTags(pos_tag))
                        clustSent = clustSent + s
                    else:    
                        finalData.append([keyW,clustSent])
                        keyW=[]
                        clustSent=""
                        rake.extract_keywords_from_text(s)
                        for x in rake.get_ranked_phrases():
                            x = x.split()
                            pos_tag = nltk.pos_tag(x)
                            keyW.append(removeTags(pos_tag))
                        clustSent = clustSent + s
        else:
            clustSent = clustSent + s
            rake.extract_keywords_from_text(s)
            for x in rake.get_ranked_phrases():
                x = x.split()
                pos_tag = nltk.pos_tag(x)
                keyW.append(removeTags(pos_tag))
    finalData.append([keyW,clustSent])
    return(finalData)

#Creating Cluster
flag=0
testFlag = 0
keywords=[]
data = ""
fullText = []
n1= []
englishStemmer=SnowballStemmer("english")
stop_words = set(stopwords.words('english')) 
for para in doc.paragraphs:
    # print(para.text)
    sentence_text = nltk.sent_tokenize(para.text)
    fullText.append(para.text)
    for i in sentence_text:
        word_token = nltk.word_tokenize(i)
        if(len(word_token) != 0):
            if(data==""):
                if(len(word_token)<4):
                    data = i + '.'
                else:
                    data = i
            else:
                if(len(word_token)<4):
                    data +=i + ". "
                else:
                    data += i + " "
cluster = clustersCreater(data)
print(cluster[0])
len(cluster)

"""# **Question Processing**"""

#Answer Prediction
question="define desertification?"
StemKey = []
Key=[]
C=''

qKeys=Rake()
qKeys.extract_keywords_from_text(question)
Key=qKeys.get_ranked_phrases()

for word in Key:
   StemKey.append(englishStemmer.stem(word))

for a in Key:
  C=C +" "+ a

print(Key)
print(StemKey)
print(C)

"""#  **STRING MATCH APPROACH FROM CORPUS**

"""

#Fetching Corpus Document
Corpus=[]
document= docx.Document('/content/drive/My Drive/SIH.docx')
for para in document.paragraphs:
    xam=sent_tokenize(para.text)
    for token in xam:
        Corpus.append(token.replace('\xa0', ' '))

#Matching Indexes with combine texts

CorpusIndex=[]
index=0

while(index<len(Corpus)):
  # print(C,Corpus[index],fuzz.token_set_ratio(C,Corpus[index]))
  if(fuzz.token_set_ratio(C,Corpus[index])>90):
        CorpusIndex.append(index)
        # print(C,Corpus[index],fuzz.token_set_ratio(C,Corpus[index]))
  index+=1

    

print(Corpus[331])
print(len(CorpusIndex))
print(CorpusIndex)

i=0
FinalString=''
appread=[]

for s in CorpusIndex:
  if s not in appread:
    i=0
    x=s
    for i in range(10):
      if(FinalString ==' ' or FinalString.endswith('.')):
        FinalString=FinalString+' '+Corpus[x]
      else:
        FinalString=FinalString+'. '+Corpus[x]
      # print(Corpus[x])
      appread.append(x)
      x+=1
# print(FinalString)
A=model_qa([FinalString],[question])
# print(A[0])
FS=sent_tokenize(FinalString)
# print(FS[3])
for fs in FS:
  if(fuzz.token_set_ratio(A[0],fs)>98):
    # print(fuzz.token_set_ratio(A[0],fs),fs,A[0])
    val=FS.index(fs)
    break

a=val
FA=''
while(a <len(FS)):
  FA=FA+''+FS[a]
  a+=1
print(question)
print(FA)

# marks=15
# SelectedStrings=[]
# para=''
# for index in Corpusindex:
#   i=1
#   while(i<=marks):
#     # print(Corpus[index])
#     para=para+' '+ Corpus[index]
#     index+=1
#     i+=1
#   SelectedStrings.append(para)
#   para=''
  
# Answer=''
# for x in SelectedStrings:
#   Answer=Answer +' '+ x

# if(marks <3):
#   for corsp in SelectedStrings:
#     ANS=model_qa([corsp],[question])
#     print(ANS[0])
# else:
#   print(Answer)
#   # print(model_qa([Answer],[question]))

"""# **Cluster**"""

print(cluster[0])
#[[[key1,key2],Paragraph]]

i=0
IndexMat=[]
while(i<len(cluster)):
  k,p=cluster[i]
  ii=0
  found=0
  
  while(ii<len(StemKey)):
    for a in k:
      if(fuzz.token_set_ratio(a,StemKey[ii])>95):
        found+=1
        break
    ii+=1
    
  if(found==len(StemKey)):
      IndexMat.append(i)
  
    
  i+=1

len(IndexMat)

ANS=[]
i=1
for x in IndexMat:
  k,p=cluster[x]
  A=(model_qa([p],[question]))
  if A not in ANS:
    ANS.append(A)
  # print(A)
  # MODEL CAN BE USED FOR GENERATING ONLY TWO/ONE MARKS QUESTION
  # ALSO FOR DETERMINIG WHERE THE ANSWER STARTS
  print(p)
  i+=1
  if(i>3):
    break